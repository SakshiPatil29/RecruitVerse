"""
Unified text extraction for uploaded files.

Consolidates what used to be a PDF-only script (src/parser/pdf_extractor.py)
into one entry point that handles the three formats the spec requires for
both resumes and job descriptions: PDF, DOCX, TXT. Resume parsing and JD
parsing both call `extract_text`, so the format-handling logic lives in
exactly one place.
"""

import io
import os

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".txt")


def _extract_pdf_text(file_bytes):
    import fitz  # PyMuPDF

    text = ""
    with fitz.open(stream=file_bytes, filetype="pdf") as pdf:
        for page in pdf:
            text += page.get_text()
    return text


def _extract_docx_text(file_bytes):
    import docx

    document = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in document.paragraphs]

    # Tables (skills/experience are sometimes laid out in a table).
    for table in document.tables:
        for row in table.rows:
            paragraphs.append(" ".join(cell.text for cell in row.cells))

    return "\n".join(paragraphs)


def _extract_txt_text(file_bytes):
    return file_bytes.decode("utf-8", errors="ignore")


def extract_text(file_bytes, filename):
    """Extract raw text from an uploaded file's bytes based on its extension.

    Raises ValueError for unsupported formats so callers can surface a
    clear message instead of a stack trace.
    """

    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        return _extract_pdf_text(file_bytes)
    if ext == ".docx":
        return _extract_docx_text(file_bytes)
    if ext == ".txt":
        return _extract_txt_text(file_bytes)

    raise ValueError(
        f"Unsupported file type '{ext}' for '{filename}'. "
        f"Supported types: {', '.join(SUPPORTED_EXTENSIONS)}"
    )


def is_supported(filename):
    return os.path.splitext(filename)[1].lower() in SUPPORTED_EXTENSIONS
